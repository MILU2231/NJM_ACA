import os
import re
import json
import numpy as np
import matplotlib

# 设置后端为 Agg，防止多进程绘图时出现 GUI 错误
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from scipy.signal import find_peaks, peak_prominences
import matplotlib.font_manager as fm
from datetime import datetime, timedelta
import warnings
from multiprocessing import Pool, cpu_count, freeze_support

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

warnings.filterwarnings('ignore')

# =========================
# 路径配置
# =========================

# 修改：读取1440优化代码生成的拼接数据路径
INPUT_DIR = r"X:\NJM_Item\ACA_对接优化数据集\原始数据拼接1440"

# 从路径中提取设备号 - 根据实际传感器信息调整
DEVICE_ID = "C13_02"  # 可以根据实际传感器信息动态获取

# 新的根目录结构 - 修改路径结构
ROOT_OUTPUT_DIR = r"X:\NJM_Item\ACA_对接优化数据集\C13_ACA列车切片数据"
DEVICE_SLICES_DIR = os.path.join(ROOT_OUTPUT_DIR, f"{DEVICE_ID}列车事件切片")

# 列车事件输出目录
OUT_PARENT_DIR = os.path.join(DEVICE_SLICES_DIR, "Min_列车事件")

# 正常振动输出目录
NORMAL_VIB_PARENT_DIR = os.path.join(DEVICE_SLICES_DIR, "Min_非列车事件")

FS = 50

# =========================
# 核心检测参数 (列车事件) - 使用去噪音2的参数
# =========================
RMS_WIN_SEC = 0.8
SMOOTH_SEC = 1.5

# 【自适应阈值参数】
# 你说正常水平在 120 左右：这里把目标靠近 120，并给一个更合理的"稳定区间"
EXPECTED_TARGET_EVENTS = 120
EXPECTED_MIN_EVENTS = 105
EXPECTED_MAX_EVENTS = 135

# 二分搜索迭代次数（更稳定）
THRESHOLD_SEARCH_ITERATIONS = 22

# 初始阈值计算方式
BASE_THRESHOLD_METHOD = 'otsu'  # 'otsu', 'percentile', 'kmeans'
FALLBACK_PERCENTILE = 98.0

EXIT_RATIO = 0.6
RECOVER_HOLD_SEC = 2

# 【抗碎裂优化：放宽持续时间范围，避免大量误丢导致 40-50 个事件】
MIN_EVENT_SEC = 10.0
MAX_EVENT_SEC = 80.0
MERGE_GAP_SEC = 5  # 原来 3，放宽一点，减少碎裂

TARGET_DURATION_SEC = 30
TARGET_SAMPLES = int(TARGET_DURATION_SEC * FS)

# 单列车事件固定截取时长
SINGLE_TRAIN_MAX_DURATION_SEC = 60
SINGLE_TRAIN_MAX_SAMPLES = int(SINGLE_TRAIN_MAX_DURATION_SEC * FS)

# 相汇检测参数
CONVERGENCE_MIN_PEAK_DISTANCE = 8
CONVERGENCE_MAX_PEAK_DISTANCE = 35
CONVERGENCE_MIN_VALLEY_RATIO = 0.4
CONVERGENCE_MIN_SECOND_PEAK_RATIO = 0.7
MIN_TRAIN_DURATION_SEC = 6

# 峰检测参数
PEAK_HEIGHT_PERCENTILE = 74
MIN_ABSOLUTE_PEAK_HEIGHT = 1.5e-5
CONVERGENCE_MIN_PEAK_PROMINENCE_RATIO = 0.15

# =========================
# 核心检测参数 (正常振动 - 无列车)
# =========================
NORMAL_SLICE_DURATION_SEC = 30  # 30s = 1500点
NORMAL_MIN_DURATION_SEC = 30
NORMAL_MARGIN_SEC = 10  # 安全边界
MAX_NORMAL_SLICES = 5000  # 最大生成数量限制提高到5000
STABILITY_THRESHOLD = 2.0  # 稳定性筛选阈值

# 调试模式
DEBUG_MODE = True


# =========================
# 字体设置
# =========================
def setup_matplotlib_fonts():
    try:
        plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']
        plt.rcParams['axes.unicode_minus'] = False
        available_fonts = set([f.name for f in fm.fontManager.ttflist])
        if 'SimHei' in available_fonts:
            plt.rcParams['font.sans-serif'] = ['SimHei', 'DejaVu Sans']
        elif 'Microsoft YaHei' in available_fonts:
            plt.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'DejaVu Sans']
        else:
            plt.rcParams['font.sans-serif'] = ['DejaVu Sans']
    except Exception as e:
        print(f"字体设置失败: {e}")
        plt.rcParams['axes.unicode_minus'] = False


setup_matplotlib_fonts()


# =========================
# 工具函数
# =========================
def ensure_dir(p: str):
    os.makedirs(p, exist_ok=True)


def nan_fill_linear(x: np.ndarray) -> np.ndarray:
    x = x.astype(np.float64)
    n = len(x)
    isn = np.isnan(x)
    if not isn.any(): return x
    idx = np.arange(n)
    valid = ~isn
    if valid.sum() == 0: raise ValueError("数据全是 NaN")
    x[isn] = np.interp(idx[isn], idx[valid], x[valid])
    return x


def moving_rms(x: np.ndarray, win: int) -> np.ndarray:
    win = max(1, int(win))
    xx = x.astype(np.float64) ** 2
    kernel = np.ones(win, dtype=np.float64) / win
    return np.sqrt(np.convolve(xx, kernel, mode="same") + 1e-18)


def moving_average(x: np.ndarray, win: int) -> np.ndarray:
    win = max(1, int(win))
    return np.convolve(x.astype(np.float64), np.ones(win, dtype=np.float64) / win, mode="same")


def merge_segments(segs, merge_gap: int):
    if not segs: return []
    segs = sorted(segs, key=lambda t: t[0])
    merged = [segs[0]]
    for s, e in segs[1:]:
        ps, pe = merged[-1]
        if s - pe <= merge_gap:
            merged[-1] = (ps, max(pe, e))
        else:
            merged.append((s, e))
    return merged


# =========================
# 自适应阈值相关函数（来自去噪音2）
# =========================
def calculate_otsu_threshold(data: np.ndarray) -> float:
    """
    使用Otsu方法计算双峰分布的最佳阈值
    """
    try:
        data = np.asarray(data, dtype=np.float64)
        data = data[np.isfinite(data)]
        if data.size == 0:
            return 0.0

        data_norm = data - np.min(data)
        if np.max(data_norm) > 0:
            data_norm = data_norm / np.max(data_norm) * 255.0
        else:
            return float(np.mean(data))

        data_int = data_norm.astype(np.uint8)

        hist, bin_edges = np.histogram(data_int, bins=256, range=(0, 256))
        bin_centers = (bin_edges[:-1] + bin_edges[1:]) / 2

        total_weight = np.sum(hist)
        if total_weight == 0:
            return float(np.mean(data))

        sum_total = np.sum(bin_centers * hist)
        sum_background = 0.0
        weight_background = 0.0
        max_variance = -1.0
        threshold = 0.0

        for i in range(256):
            weight_background += hist[i]
            if weight_background == 0:
                continue

            weight_foreground = total_weight - weight_background
            if weight_foreground == 0:
                break

            sum_background += bin_centers[i] * hist[i]
            mean_background = sum_background / weight_background
            mean_foreground = (sum_total - sum_background) / weight_foreground

            variance = weight_background * weight_foreground * (mean_background - mean_foreground) ** 2
            if variance > max_variance:
                max_variance = variance
                threshold = bin_centers[i]

        threshold_original = threshold / 255.0 * (np.max(data) - np.min(data)) + np.min(data)
        return float(threshold_original)

    except Exception as e:
        print(f"   Otsu阈值计算失败: {e}，使用均值+标准差方法")
        data = np.asarray(data, dtype=np.float64)
        data = data[np.isfinite(data)]
        if data.size == 0:
            return 0.0
        return float(np.mean(data) + 2 * np.std(data))


def calculate_adaptive_threshold(rms_s: np.ndarray, method: str = 'otsu') -> float:
    """
    计算初始阈值（注意：这里不要用过强的"最低阈值夹死"，否则容易导致事件数偏少）
    """
    rms_s = np.asarray(rms_s, dtype=np.float64)
    rms_s = rms_s[np.isfinite(rms_s)]
    if rms_s.size == 0:
        return 0.0

    if method == 'otsu':
        threshold = calculate_otsu_threshold(rms_s)
    elif method == 'percentile':
        threshold = float(np.nanpercentile(rms_s, FALLBACK_PERCENTILE))
    elif method == 'kmeans':
        try:
            sorted_data = np.sort(rms_s)
            split_idx = int(len(sorted_data) * 0.95)
            threshold = float(sorted_data[split_idx])
        except Exception:
            threshold = float(np.nanpercentile(rms_s, FALLBACK_PERCENTILE))
    else:
        threshold = float(np.mean(rms_s) + 2.5 * np.std(rms_s))

    # 更温和的下限：用中位数附近的能量水平做下限，避免把阈值抬得过高
    floor_thr = float(max(1e-7, np.nanpercentile(rms_s, 60)))
    threshold = max(threshold, floor_thr)
    return float(threshold)


def detect_events_with_threshold(rms_s: np.ndarray, fs: int, enter_thr: float):
    """
    使用给定阈值检测事件（enter/exit）
    """
    exit_thr = enter_thr * EXIT_RATIO
    hold = int(RECOVER_HOLD_SEC * fs)

    in_event, start, low_count, segs = False, None, 0, []
    for i, v in enumerate(rms_s):
        if not in_event:
            if v > enter_thr:
                in_event, start, low_count = True, i, 0
        else:
            if v < exit_thr:
                low_count += 1
                if low_count >= hold:
                    segs.append((start, i - hold + 1))
                    in_event = False
            else:
                low_count = 0

    if in_event and start is not None:
        segs.append((start, len(rms_s)))

    return segs, enter_thr, exit_thr


def _filter_and_merge_segments(segs, fs: int):
    min_event_samples = int(MIN_EVENT_SEC * fs)
    max_event_samples = int(MAX_EVENT_SEC * fs)
    filtered = [(s, e) for s, e in segs if min_event_samples <= (e - s) <= max_event_samples]
    merged = merge_segments(filtered, int(MERGE_GAP_SEC * fs))
    return merged


def detect_events_adaptive(rms_s: np.ndarray, fs: int):
    """
    【核心改进】更稳定的自适应阈值：
    - 先构造阈值上下界（夹逼）
    - 再二分搜索，让事件数稳定逼近 120（或你的区间）
    """
    print(f"\n🔍 开始自适应阈值检测 (夹逼+二分)...")

    rms_s = np.asarray(rms_s, dtype=np.float64)
    rms_s = rms_s[np.isfinite(rms_s)]
    if rms_s.size == 0:
        print("   ⚠️ RMS数据为空，返回0事件")
        return [], 0.0, 0.0

    initial = calculate_adaptive_threshold(rms_s, method=BASE_THRESHOLD_METHOD)
    print(f"   初始阈值 ({BASE_THRESHOLD_METHOD}方法): {initial:.6f}")

    def count_events(thr: float):
        segs, _, _ = detect_events_with_threshold(rms_s, fs, thr)
        segs2 = _filter_and_merge_segments(segs, fs)
        return len(segs2), segs2

    # 1) 先用分位数构造一个比较"稳"的搜索区间
    # 低阈值(事件多) / 高阈值(事件少)
    low = float(np.nanpercentile(rms_s, 55))
    high = float(np.nanpercentile(rms_s, 99.5))
    low = min(low, initial)
    high = max(high, initial)

    # 防止 low==high
    if not (high > low):
        high = low * 1.5 + 1e-12

    # 2) 夹逼：保证 low 事件数 >= 目标（或至少不太少），high 事件数 <= 目标（或至少不太多）
    #    事件数对阈值通常单调：阈值越大，事件越少
    c_low, seg_low = count_events(low)
    c_high, seg_high = count_events(high)

    # 如果 low 仍然太少 -> 继续降低 low
    expand_step = 0
    while c_low < EXPECTED_TARGET_EVENTS and expand_step < 12:
        low = max(1e-7, low / 1.6)
        c_low, seg_low = count_events(low)
        expand_step += 1

    # 如果 high 仍然太多 -> 继续提高 high
    expand_step = 0
    while c_high > EXPECTED_TARGET_EVENTS and expand_step < 12:
        high = high * 1.6
        c_high, seg_high = count_events(high)
        expand_step += 1

    print(f"   夹逼区间: low={low:.6f}(事件={c_low}), high={high:.6f}(事件={c_high})")

    # 如果夹逼失败（比如怎么调都上不去/下不来），就用最接近目标的一个
    best_thr = initial
    best_segs = []
    best_cnt = 0
    for thr_try, seg_try, cnt_try in [(low, seg_low, c_low), (high, seg_high, c_high), (initial, None, None)]:
        if thr_try == initial:
            cnt_try, seg_try = count_events(initial)
        if best_cnt == 0 or abs(cnt_try - EXPECTED_TARGET_EVENTS) < abs(best_cnt - EXPECTED_TARGET_EVENTS):
            best_thr, best_segs, best_cnt = thr_try, seg_try, cnt_try

    # 3) 二分搜索
    for it in range(THRESHOLD_SEARCH_ITERATIONS):
        mid = (low + high) / 2.0
        cnt, segs_mid = count_events(mid)

        print(f"   二分 {it + 1:02d}: 阈值={mid:.6f}, 事件数={cnt}")

        if abs(cnt - EXPECTED_TARGET_EVENTS) < abs(best_cnt - EXPECTED_TARGET_EVENTS):
            best_thr, best_segs, best_cnt = mid, segs_mid, cnt

        if EXPECTED_MIN_EVENTS <= cnt <= EXPECTED_MAX_EVENTS:
            print(f"   ✅ 命中区间: 阈值={mid:.6f}, 事件数={cnt}")
            return segs_mid, mid, mid * EXIT_RATIO

        # cnt 太多 -> 阈值偏低 -> 提高阈值 -> 移动 low
        if cnt > EXPECTED_MAX_EVENTS:
            low = mid
        # cnt 太少 -> 阈值偏高 -> 降低阈值 -> 移动 high
        elif cnt < EXPECTED_MIN_EVENTS:
            high = mid
        else:
            break

    print(f"   ⚠️ 二分结束，使用最佳阈值: {best_thr:.6f}, 事件数: {best_cnt}")
    return best_segs, best_thr, best_thr * EXIT_RATIO


def save_json(output_dir, filename, data_dict):
    path = os.path.join(output_dir, filename)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data_dict, f, indent=4, ensure_ascii=False)


def extract_date_from_path(file_path):
    """
    从文件路径中提取日期信息
    假设路径中包含 YYYY-MM-DD 格式的日期
    """
    date_pattern = r'(\d{4}-\d{2}-\d{2})'
    match = re.search(date_pattern, file_path)
    if match:
        return match.group(1)
    else:
        # 如果没有找到日期，使用当前日期
        return datetime.now().strftime("%Y-%m-%d")


def calculate_timestamps(start_date, start_index, end_index, fs):
    """
    计算时间戳
    """
    try:
        # 假设数据从当天的 00:00:00 开始
        base_datetime = datetime.strptime(start_date, "%Y-%m-%d")

        # 计算时间偏移
        start_time_offset = start_index / fs
        end_time_offset = end_index / fs
        duration = (end_index - start_index) / fs

        # 计算完整时间戳
        start_timestamp = base_datetime + timedelta(seconds=start_time_offset)
        end_timestamp = base_datetime + timedelta(seconds=end_time_offset)

        # 格式化为字符串
        start_time_str = start_timestamp.strftime("%Y-%m-%d %H:%M:%S")
        end_time_str = end_timestamp.strftime("%Y-%m-%d %H:%M:%S")

        return {
            "start_time": start_time_str,
            "end_time": end_time_str,
            "duration_seconds": float(duration),
            "start_index": int(start_index),
            "end_index": int(end_index)
        }
    except Exception as e:
        print(f"时间戳计算失败: {e}")
        return {
            "start_time": "未知",
            "end_time": "未知",
            "duration_seconds": float((end_index - start_index) / fs),
            "start_index": int(start_index),
            "end_index": int(end_index)
        }


def generate_time_series(start_time_str, num_points, fs):
    """
    生成时间序列，用于原始数据图的x轴
    """
    try:
        base_time = datetime.strptime(start_time_str, "%Y-%m-%d %H:%M:%S")
        time_series = [base_time + timedelta(seconds=i / fs) for i in range(num_points)]
        return time_series
    except:
        # 如果失败，返回相对时间
        return np.arange(num_points) / fs


# =========================
# PART 1: 正常振动处理逻辑 (使用去噪音2的逻辑)
# =========================
def extract_normal_vibration_segments(events, total_length, fs):
    if not events: return [(0, total_length)]
    norms = []
    if events[0][0] > 0: norms.append((0, events[0][0]))
    for i in range(len(events) - 1):
        if events[i + 1][0] > events[i][1]: norms.append((events[i][1], events[i + 1][0]))
    if events[-1][1] < total_length: norms.append((events[-1][1], total_length))
    m, min_s = int(NORMAL_MARGIN_SEC * fs), int(NORMAL_MIN_DURATION_SEC * fs)
    res = []
    for s, e in norms:
        ss, ee = s + m, e - m
        if ee - ss >= min_s: res.append((ss, ee))
    return res


def check_slice_stability(slice_data, threshold=STABILITY_THRESHOLD):
    """稳定性检测：拒绝前小后大或波动不均的数据"""
    n = len(slice_data)
    chunks = np.array_split(slice_data, 3)
    stds = [np.std(chunk) for chunk in chunks]
    max_std = max(stds)
    min_std = min(stds)
    if min_std < 1e-9:
        if max_std < 1e-6: return True
        return False
    return (max_std / min_std) <= threshold


def process_normal_vibration_slices(x, events, fs, run_dir, start_date):
    """处理正常振动切片"""
    total_length = len(x)
    normal_segments = extract_normal_vibration_segments(events, total_length, fs)
    slice_dur = int(NORMAL_SLICE_DURATION_SEC * fs)

    candidate_slices = []
    for ss, ee in normal_segments:
        num = (ee - ss) // slice_dur
        for i in range(num):
            candidate_slices.append((ss + i * slice_dur, ss + (i + 1) * slice_dur))

    if not candidate_slices: return None

    # 修改：移除npy目录，只保留json和图片目录
    img_dir = os.path.join(run_dir, "images")
    json_dir = os.path.join(run_dir, "json_slices")
    for d in [img_dir, json_dir]: ensure_dir(d)

    doc = Document()
    doc.add_heading(f"正常振动切片分析 (Top {MAX_NORMAL_SLICES})", level=1)

    # 添加时间信息到Word文档
    time_info_para = doc.add_paragraph()
    time_info_para.add_run(f"数据日期: {start_date} | 采样率: {fs} Hz | 切片时长: {NORMAL_SLICE_DURATION_SEC}秒\n")

    saved_count = 0
    discarded_count = 0

    for i, (s, e) in enumerate(candidate_slices, start=1):
        if saved_count >= MAX_NORMAL_SLICES:
            break

        slice_data = x[s:e]
        if not check_slice_stability(slice_data):
            discarded_count += 1
            continue

        saved_count += 1

        # 计算时间戳
        time_info = calculate_timestamps(start_date, s, e, fs)

        # 保存 JSON - 按照要求顺序排列，添加label字段
        json_data = {
            "start_time": time_info["start_time"],
            "end_time": time_info["end_time"],
            "duration_seconds": time_info["duration_seconds"],
            "start_index": time_info["start_index"],
            "end_index": time_info["end_index"],
            "event_id": saved_count,
            "train_index": saved_count,
            "type": "normal_vibration",
            "is_borrowed": False,
            "borrowed_length": 0,
            "sampling_rate": fs,
            "label": 0,  # 添加label字段，正常振动label=0
            "data": slice_data.tolist()
        }
        save_json(json_dir, f"normal_{saved_count:04d}.json", json_data)

        # ============================================
        # 暂时注释掉非列车事件的图片保存 - 修改部分（保留原注释）
        # ============================================
        # 绘图 - 原始数据使用时间戳作为x轴
        # plt.figure(figsize=(10, 3.4))
        # time_series = generate_time_series(time_info["start_time"], len(slice_data), fs)
        # if isinstance(time_series[0], datetime):
        #     plt.plot(time_series, slice_data, color='blue', alpha=0.7, linewidth=0.8)
        #     plt.gcf().autofmt_xdate()  # 自动格式化日期
        # else:
        #     plt.plot(np.arange(s, e) / fs, slice_data, color='blue', alpha=0.7, linewidth=0.8)
        # plt.title(f"正常片段 {saved_count:04d} | Std: {np.std(slice_data):.2e}")
        # plt.xlabel(f"时间 (开始: {time_info['start_time']})")
        # plt.tight_layout()
        # img_path = os.path.join(img_dir, f"normal_{saved_count:04d}.png")
        # plt.savefig(img_path, dpi=180);
        # plt.close()
        # ============================================

        # Word文档添加详细信息
        doc.add_heading(f"片段 {saved_count:04d}", level=2)

        # 时间信息段落
        time_para = doc.add_paragraph()
        time_para.add_run(f"开始时间: {time_info['start_time']}\n").bold = True
        time_para.add_run(f"结束时间: {time_info['end_time']}\n")
        time_para.add_run(f"持续时间: {time_info['duration_seconds']:.2f} 秒\n")
        time_para.add_run(f"标签: 0 (非列车事件)\n")  # 添加标签信息
        time_para.add_run(f"原始索引: {s} - {e}\n")

        # ============================================
        # 暂时注释掉图片插入Word文档 - 修改部分（保留原注释）
        # ============================================
        # doc.add_picture(img_path, width=Inches(6.0))
        doc.add_paragraph()  # 空行分隔

    # 添加统计信息
    stats_para = doc.add_paragraph()
    stats_para.add_run(f"\n统计摘要:\n").bold = True
    stats_para.add_run(f"总生成切片: {saved_count}\n")
    stats_para.add_run(f"丢弃切片: {discarded_count} (稳定性不合格)\n")
    stats_para.add_run(f"数据日期: {start_date}\n")
    stats_para.add_run(f"标签: 0 (非列车事件)\n")  # 添加标签信息

    doc.save(os.path.join(run_dir, "normal_vibration_report.docx"))
    return {"total": saved_count, "discarded": discarded_count}


# =========================
# PART 2: 列车事件处理逻辑 (使用去噪音2的双峰检测方法)
# =========================
def generate_slices(coarse_segs, n, fs):
    slices = []
    for i, (start, end) in enumerate(coarse_segs):
        base_end = min(end, start + SINGLE_TRAIN_MAX_SAMPLES)
        if i + 1 < len(coarse_segs):
            next_start = coarse_segs[i + 1][0]
            safe_end = min(base_end, next_start - int(2.0 * fs))
        else:
            safe_end = base_end
        final_end = min(safe_end, n)
        if final_end - start < int(MIN_EVENT_SEC * fs): continue
        slices.append((start, final_end))
    return slices


def standardize_length_with_peak(signal: np.ndarray, target_length: int, fs: int):
    if signal is None or len(signal) == 0: return np.zeros(target_length)
    current_length = len(signal)
    if current_length == target_length:
        return signal
    elif current_length > target_length:
        try:
            smooth_win = min(50, current_length // 20)
            smoothed = np.convolve(signal, np.ones(smooth_win) / smooth_win, mode='same') if smooth_win > 3 else signal
            peaks, properties = find_peaks(smoothed, height=np.percentile(smoothed, 70), distance=int(5 * fs),
                                           prominence=np.std(smoothed) * 0.2)
            if len(peaks) > 0:
                main_peak_idx = peaks[np.argmax(properties['peak_heights'])]
                start_idx = max(0, min(main_peak_idx - target_length // 2, current_length - target_length))
                return signal[int(start_idx):int(start_idx + target_length)]
        except:
            pass
        start_idx = max(0, (current_length - target_length) // 2)
        return signal[int(start_idx):int(start_idx + target_length)]
    else:
        pad_length = target_length - current_length
        return np.pad(signal, (0, pad_length), mode='reflect')


def standardize_convergence_train(train_data: np.ndarray, peak_idx_in_train: int, target_length: int, fs: int):
    # 与原代码逻辑一致，处理相汇事件的标准化
    if train_data is None or len(train_data) == 0: return np.zeros(target_length)
    current_length = len(train_data)
    if current_length >= target_length:
        start_idx = max(0, min(peak_idx_in_train - target_length // 2, current_length - target_length))
        return train_data[int(start_idx):int(start_idx + target_length)]
    else:
        return np.pad(train_data, (0, target_length - current_length), mode='reflect')


def process_single_train_event(original_data, fs, target_samples, max_duration_samples):
    if len(original_data) > max_duration_samples:
        truncated = original_data[:max_duration_samples]
    else:
        truncated = original_data
    return standardize_length_with_peak(truncated, target_samples, fs)


def is_convergence_event(signal: np.ndarray, fs: int):
    if len(signal) < 100: return False, []
    try:
        smooth_win = min(50, len(signal) // 20)
        smoothed = np.convolve(signal, np.ones(smooth_win) / smooth_win,
                               mode='same') if smooth_win > 3 else signal.copy()

        signal_std = np.std(smoothed)
        height_threshold = max(np.percentile(np.abs(smoothed), PEAK_HEIGHT_PERCENTILE),
                               np.mean(smoothed) + signal_std * 0.5, MIN_ABSOLUTE_PEAK_HEIGHT)

        peaks, props = find_peaks(np.abs(smoothed), height=height_threshold, distance=int(3 * fs),
                                  prominence=signal_std * CONVERGENCE_MIN_PEAK_PROMINENCE_RATIO)

        if len(peaks) < 2: return False, []

        # 取最高的两个峰
        sorted_indices = np.argsort(props['peak_heights'])[::-1]
        top_peaks = np.sort(peaks[sorted_indices[:2]])
        peak1, peak2 = top_peaks

        # 修正高度获取方式，因为排序后索引变了
        idx1 = np.where(peaks == peak1)[0][0]
        idx2 = np.where(peaks == peak2)[0][0]
        peak1_height = props['peak_heights'][idx1]
        peak2_height = props['peak_heights'][idx2]

        peak_distance = (peak2 - peak1) / fs
        if not (CONVERGENCE_MIN_PEAK_DISTANCE <= peak_distance <= CONVERGENCE_MAX_PEAK_DISTANCE): return False, []

        valley_signal = smoothed[peak1:peak2]
        if len(valley_signal) == 0: return False, []
        valley_ratio = np.min(np.abs(valley_signal)) / peak1_height if peak1_height > 0 else 0
        if valley_ratio > CONVERGENCE_MIN_VALLEY_RATIO: return False, []

        second_peak_ratio = peak2_height / peak1_height if peak1_height > 0 else 0
        if second_peak_ratio < CONVERGENCE_MIN_SECOND_PEAK_RATIO: return False, []

        train_segments = []
        exit_threshold = np.mean(smoothed) + signal_std * 0.1

        for peak_idx in [peak1, peak2]:
            start_idx, end_idx = peak_idx, peak_idx
            # 向前搜索
            for j in range(min(peak_idx, int(15 * fs))):
                if abs(smoothed[peak_idx - j]) < exit_threshold:
                    if j > 2: start_idx = peak_idx - j; break
            # 向后搜索
            for j in range(min(len(smoothed) - peak_idx - 1, int(20 * fs))):
                if abs(smoothed[peak_idx + j]) < exit_threshold:
                    if j > 4: end_idx = peak_idx + j; break

            duration = (end_idx - start_idx) / fs
            if duration < MIN_TRAIN_DURATION_SEC: continue
            train_segments.append(
                {'peak_index': int(peak_idx), 'start_index': int(start_idx), 'end_index': int(end_idx),
                 'duration': duration, 'peak_idx_in_train': int(peak_idx - start_idx)})

        if len(train_segments) >= 2:
            train_segments.sort(key=lambda x: x['start_index'])
            # 简单重叠处理
            if train_segments[1]['start_index'] < train_segments[0]['end_index']:
                mid = (train_segments[0]['end_index'] + train_segments[1]['start_index']) // 2
                train_segments[0]['end_index'] = mid
                train_segments[1]['start_index'] = mid
                train_segments[0]['peak_idx_in_train'] = train_segments[0]['peak_index'] - train_segments[0][
                    'start_index']
                train_segments[1]['peak_idx_in_train'] = train_segments[1]['peak_index'] - train_segments[1][
                    'start_index']
            return True, train_segments

        return False, []
    except:
        return False, []


def process_convergence_event(event_signal, train_segments, fs, target_samples):
    separated = []
    event_len = len(event_signal)

    # Train 1
    t1 = train_segments[0]
    data1 = event_signal[t1['start_index']:t1['end_index']]

    # Borrow logic
    is_borrowed, borrow_len = False, 0
    if len(data1) < target_samples:
        borrow_start = train_segments[1]['end_index']
        if borrow_start < event_len:
            need = target_samples - len(data1)
            avail = min(need, int(10 * fs), event_len - borrow_start)
            if avail > 0 and np.std(event_signal[borrow_start:borrow_start + avail]) > 1e-6:
                data1 = np.concatenate([data1, event_signal[borrow_start:borrow_start + avail]])
                is_borrowed, borrow_len = True, avail

    std1 = standardize_convergence_train(data1, t1['peak_idx_in_train'], target_samples, fs)
    separated.append({'data': std1, 'is_borrowed': is_borrowed, 'borrowed_length': borrow_len})

    # Train 2
    t2 = train_segments[1]
    data2 = event_signal[t2['start_index']:t2['end_index']]
    std2 = standardize_convergence_train(data2, t2['peak_idx_in_train'], target_samples, fs)
    separated.append({'data': std2, 'is_borrowed': False, 'borrowed_length': 0})

    return separated


def create_event_visualization(event_id, original, is_conv, train_segs, separated, fs, save_path, time_info=None):
    plt.close('all')
    try:
        if is_conv and len(train_segs) >= 2:
            fig, axes = plt.subplots(3, 1, figsize=(12, 9))

            # 原始数据图 - 使用时间戳作为x轴
            if time_info:
                time_series = generate_time_series(time_info["start_time"], len(original), fs)
                if isinstance(time_series[0], datetime):
                    axes[0].plot(time_series, original, 'b-', lw=1)
                    axes[0].tick_params(axis='x', rotation=45)
                else:
                    axes[0].plot(np.arange(len(original)) / fs, original, 'b-', lw=1)
            else:
                axes[0].plot(np.arange(len(original)) / fs, original, 'b-', lw=1)

            axes[0].set_title(f"事件 {event_id} - 原始相汇信号 | 标签: 1")  # 添加标签信息
            if time_info:
                axes[0].set_xlabel(
                    f"时间 | 开始: {time_info['start_time']} | 时长: {time_info['duration_seconds']:.1f}s")

            for i, seg in enumerate(train_segs[:2]):
                color = ['red', 'green'][i]
                if time_info and isinstance(time_series[0], datetime):
                    seg_start_time = time_series[0] + timedelta(seconds=seg['start_index'] / fs)
                    seg_end_time = time_series[0] + timedelta(seconds=seg['end_index'] / fs)
                    axes[0].axvspan(seg_start_time, seg_end_time, color=color, alpha=0.15)
                    peak_time = time_series[0] + timedelta(seconds=seg['peak_index'] / fs)
                    axes[0].plot(peak_time, original[seg['peak_index']], 'o', color=color)
                else:
                    axes[0].axvspan(seg['start_index'] / fs, seg['end_index'] / fs, color=color, alpha=0.15)
                    axes[0].plot(seg['peak_index'] / fs, original[seg['peak_index']], 'o', color=color)

            # 标准化后的数据图 - 使用相对时间
            if len(separated) > 0:
                axes[1].plot(np.arange(len(separated[0]['data'])) / fs, separated[0]['data'], 'r-')
                axes[1].set_title("分离列车1" + (" (Borrowed)" if separated[0]['is_borrowed'] else "") + " | 标签: 1")
                axes[1].set_xlabel("相对时间 (s)")
            if len(separated) > 1:
                axes[2].plot(np.arange(len(separated[1]['data'])) / fs, separated[1]['data'], 'g-')
                axes[2].set_title("分离列车2 | 标签: 1")
                axes[2].set_xlabel("相对时间 (s)")
        else:
            fig, axes = plt.subplots(2, 1, figsize=(12, 6))

            # 原始数据图 - 使用时间戳作为x轴
            if time_info:
                time_series = generate_time_series(time_info["start_time"], len(original), fs)
                if isinstance(time_series[0], datetime):
                    axes[0].plot(time_series, original, 'b-', lw=1)
                    axes[0].tick_params(axis='x', rotation=45)
                else:
                    axes[0].plot(np.arange(len(original)) / fs, original, 'b-', lw=1)
            else:
                axes[0].plot(np.arange(len(original)) / fs, original, 'b-', lw=1)

            axes[0].set_title(f"事件 {event_id} - 原始单列车信号 | 标签: 1")  # 添加标签信息
            if time_info:
                axes[0].set_xlabel(
                    f"时间 | 开始: {time_info['start_time']} | 时长: {time_info['duration_seconds']:.1f}s")

            # 标准化后的数据图 - 使用相对时间
            if separated:
                axes[1].plot(np.arange(len(separated[0]['data'])) / fs, separated[0]['data'], 'b-')
                axes[1].set_title("标准化单列车 | 标签: 1")  # 添加标签信息
                axes[1].set_xlabel("相对时间 (s)")

        plt.tight_layout()
        plt.savefig(save_path, dpi=180)
        plt.close(fig)
    except Exception as e:
        print(f"创建可视化失败 {event_id}: {e}")


def create_word_report(run_dir, stats, convergence_details, all_trains, proc_time, img_dir, conv_img_dir, start_date,
                       npy_filename, threshold_info):
    doc = Document()
    doc.add_heading('列车过桥事件处理报告', 0)

    # 添加时间信息
    time_para = doc.add_paragraph()
    time_para.add_run(
        f"数据日期: {start_date} | 处理时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").bold = True
    time_para.add_run(f"采样率: {FS} Hz | 目标时长: {TARGET_DURATION_SEC}秒\n")

    doc.add_heading('统计摘要', 1)
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Light Grid'
    data = [
        ("总列车数", str(stats['total_trains'])),
        ("单列车事件", str(stats['single_events'])),
        ("相汇事件", str(stats['convergence_events'])),
        ("分离列车数", str(stats['trains_from_convergence'])),
        ("借用数据列车", str(stats['borrowed_trains'])),
        ("错误事件", str(stats['error_events'])),
        ("自适应阈值", f"{threshold_info['enter_threshold']:.6f}"),
        ("处理耗时", f"{proc_time:.2f}s")
    ]
    for i, (k, v) in enumerate(data):
        table.cell(i, 0).text = k;
        table.cell(i, 1).text = v

    # 添加时间戳信息章节
    doc.add_heading('时间信息说明', 1)
    time_info_para = doc.add_paragraph()
    time_info_para.add_run("时间戳说明:\n").bold = True
    time_info_para.add_run(f"• 数据基准日期: {start_date}\n")
    time_info_para.add_run("• 假设数据从当天 00:00:00 开始记录\n")
    time_info_para.add_run("• 所有时间戳基于采样率50Hz计算\n")
    time_info_para.add_run("• 时间格式: YYYY-MM-DD HH:MM:SS\n")
    time_info_para.add_run("• 原始数据图横坐标为绝对时间，标准化数据图横坐标为相对时间\n")

    doc.add_heading('标签信息', 1)
    label_para = doc.add_paragraph()
    label_para.add_run("标签说明:\n").bold = True
    label_para.add_run("• 列车事件 (列车过桥): label = 1\n")
    label_para.add_run("• 正常振动 (非列车): label = 0\n")
    label_para.add_run("• 所有切片都包含label字段\n")

    doc.add_heading('事件可视化选录', 1)

    # 单列车图
    if os.path.exists(img_dir):
        doc.add_heading('单列车事件 (Top 10)', 2)
        img_files = sorted([x for x in os.listdir(img_dir) if x.endswith('.png')])
        for f in img_files[:10]:
            try:
                doc.add_picture(os.path.join(img_dir, f), width=Inches(6.0))
                # 从JSON获取时间信息
                event_id = int(re.search(r'event_(\d+)', f).group(1))
                json_file = f"event_{event_id:03d}_single.json"
                json_path = os.path.join(run_dir, "slices_json", json_file)
                if os.path.exists(json_path):
                    with open(json_path, 'r', encoding='utf-8') as jf:
                        event_data = json.load(jf)
                        time_str = f"时间: {event_data.get('start_time', '未知')} - {event_data.get('end_time', '未知')} | 时长: {event_data.get('duration_seconds', 0):.1f}s | 标签: {event_data.get('label', 1)}"
                        doc.add_paragraph(time_str)
                doc.add_paragraph(f)
            except Exception as e:
                print(f"添加单列车图片失败 {f}: {e}")

    # 相汇图
    if os.path.exists(conv_img_dir):
        doc.add_heading('相汇事件 (Top 10)', 2)
        img_files = sorted([x for x in os.listdir(conv_img_dir) if x.endswith('.png')])
        for f in img_files[:10]:
            try:
                doc.add_picture(os.path.join(conv_img_dir, f), width=Inches(6.0))
                # 从JSON获取时间信息
                event_id = int(re.search(r'event_(\d+)', f).group(1))
                json_file = f"event_{event_id:03d}_conv_train1.json"
                json_path = os.path.join(run_dir, "slices_json", json_file)
                if os.path.exists(json_path):
                    with open(json_path, 'r', encoding='utf-8') as jf:
                        event_data = json.load(jf)
                        time_str = f"时间: {event_data.get('start_time', '未知')} - {event_data.get('end_time', '未知')} | 时长: {event_data.get('duration_seconds', 0):.1f}s | 标签: {event_data.get('label', 1)}"
                        doc.add_paragraph(time_str)
                doc.add_paragraph(f)
            except Exception as e:
                print(f"添加相汇图片失败 {f}: {e}")

    # 添加详细事件列表
    doc.add_heading('事件详细信息', 1)
    doc.add_paragraph("所有事件的详细时间信息可在对应的JSON文件中查看。")

    doc.save(os.path.join(run_dir, "处理报告.docx"))


# =========================
# 读取1440拼接数据的函数
# =========================
def load_concatenated_json_data(json_path):
    """
    读取1440拼接代码生成的JSON数据文件

    Parameters:
    json_path: JSON文件路径

    Returns:
    tuple: (data_array, sensor_id, sensor_name, date_str)
    """
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            json_data = json.load(f)

        # 提取数据
        data = np.array(json_data.get("data", []))
        sensor_id = json_data.get("sensor_id", "未知")
        sensor_name = json_data.get("sensor_name", "未知传感器")
        date_str = json_data.get("date", "")

        return data, sensor_id, sensor_name, date_str
    except Exception as e:
        print(f"读取JSON文件失败 {json_path}: {e}")
        return None, None, None, None


def find_json_files(input_dir):
    """
    在输入目录中查找所有JSON文件

    Parameters:
    input_dir: 输入目录路径

    Returns:
    list: JSON文件路径列表
    """
    json_files = []

    # 递归搜索所有JSON文件
    for root, dirs, files in os.walk(input_dir):
        for file in files:
            if file.endswith('.json') and file.startswith('20') and '_concatenated.json' in file:
                full_path = os.path.join(root, file)
                json_files.append(full_path)

    return json_files


# =========================
# 标签工具函数（来自去噪音2）
# =========================
def create_labeled_npy(data, label, save_path):
    labeled_data = {'data': data, 'label': label, 'timestamp': datetime.now().isoformat()}
    np.save(save_path, labeled_data)
    return labeled_data


def save_labeled_json(output_dir, filename, data_dict, label):
    data_dict['label'] = int(label)
    save_json(output_dir, filename, data_dict)


# =========================
# 处理单个文件的逻辑
# =========================
def process_single_file(json_path):
    # 每个进程初始化自己的 Matplotlib 字体配置
    setup_matplotlib_fonts()

    file_basename = os.path.basename(json_path)
    print(f"\n{'=' * 80}")
    print(f"开始处理文件: {file_basename}")

    start_time = datetime.now()

    # 1. 加载1440拼接的JSON数据
    data, sensor_id, sensor_name, start_date = load_concatenated_json_data(json_path)

    if data is None or len(data) == 0:
        return {"file": file_basename, "status": "error", "reason": "数据加载失败"}

    # 2. 从文件名中提取日期
    if not start_date:
        start_date = extract_date_from_path(json_path)

    # 3. 确保新目录结构存在
    ensure_dir(ROOT_OUTPUT_DIR)
    ensure_dir(DEVICE_SLICES_DIR)
    ensure_dir(OUT_PARENT_DIR)
    ensure_dir(NORMAL_VIB_PARENT_DIR)

    # 4. 创建具体的运行目录 - 直接使用日期目录
    event_run_dir = os.path.join(OUT_PARENT_DIR, start_date)
    normal_run_dir = os.path.join(NORMAL_VIB_PARENT_DIR, start_date)

    # 确保目录存在
    ensure_dir(event_run_dir)
    ensure_dir(normal_run_dir)

    # 列车事件目录结构
    evt_img_dir = os.path.join(event_run_dir, "images")
    evt_conv_img_dir = os.path.join(event_run_dir, "convergence_images")
    evt_json_dir = os.path.join(event_run_dir, "slices_json")
    for d in [evt_img_dir, evt_conv_img_dir, evt_json_dir]: ensure_dir(d)

    print(f"📊 数据加载: {file_basename}")
    print(f"传感器: {sensor_id} - {sensor_name}")
    print(f"日期: {start_date}")
    print(f"数据总量: {len(data)} 点 ({len(data) / FS:.2f}s)")

    # 检查数据是否有NaN值
    if np.isnan(data).any():
        data = nan_fill_linear(data)
    x = data.astype(np.float64)

    # 5. 去趋势
    median_val = np.nanmedian(x)
    print(f"   执行基线去偏: 中位数偏移 {median_val:.2e}")
    x = x - median_val

    # 6. 自适应阈值检测（使用去噪音2的方法）
    print(f"\n🔍 开始自适应阈值检测 (夹逼+二分)...")
    rms_s = moving_average(moving_rms(x, int(RMS_WIN_SEC * FS)), int(SMOOTH_SEC * FS))
    filtered_segs, enter_thr, exit_thr = detect_events_adaptive(rms_s, FS)

    threshold_info = {
        'enter_threshold': float(enter_thr),
        'exit_threshold': float(exit_thr),
        'method': BASE_THRESHOLD_METHOD
    }

    print(f"检测到列车事件段: {len(filtered_segs)} 个")

    # ==========================================
    # 7. 处理正常振动
    # ==========================================
    print(f"正在处理正常振动: {file_basename}")
    norm_stats = process_normal_vibration_slices(x, filtered_segs, FS, normal_run_dir, start_date)
    if norm_stats:
        print(f"正常振动处理完成: 生成 {norm_stats['total']} 个切片, 丢弃 {norm_stats['discarded']} 个")
    else:
        print("正常振动处理完成: 无有效切片")

    # ==========================================
    # 8. 处理列车事件
    # ==========================================
    print(f"正在处理列车事件: {file_basename}")

    # 生成合理的列车事件切片
    slices = generate_slices(filtered_segs, len(x), FS)
    print(f"生成列车事件切片: {len(slices)} 个")

    all_processed_trains = []
    conv_stats = {'count': 0, 'trains': 0, 'borrowed': 0}
    single_count = 0
    conv_details = []
    error_events = []

    for i, (start_idx, end_idx) in enumerate(slices, 1):
        try:
            original_data = x[start_idx:end_idx]
            if len(original_data) == 0: continue

            # 计算原始事件的时间戳信息
            original_time_info = calculate_timestamps(start_date, start_idx, end_idx, FS)
            is_conv, train_segments = is_convergence_event(original_data, FS)

            if is_conv and len(train_segments) >= 2:
                # === 相汇事件处理 ===
                conv_stats['count'] += 1
                separated = process_convergence_event(original_data, train_segments, FS, TARGET_SAMPLES)
                conv_stats['trains'] += len(separated)
                has_borrowed = any(t['is_borrowed'] for t in separated)
                if has_borrowed: conv_stats['borrowed'] += 1

                conv_details.append(
                    {'event_id': i, 'train_count': len(train_segments), 'separated_count': len(separated),
                     'has_borrowed': has_borrowed})

                # 绘图 - 传入时间信息
                vis_path = os.path.join(evt_conv_img_dir, f"convergence_event_{i:03d}.png")
                create_event_visualization(i, original_data, True, train_segments, separated, FS, vis_path,
                                           original_time_info)

                # 保存数据 - 只保存JSON
                for j, train_info in enumerate(separated):
                    t_data = train_info['data']
                    fname_base = f"event_{i:03d}_conv_train{j + 1}"

                    # 计算列车在原始数据中的绝对索引
                    train_abs_start = start_idx + train_segments[j]['start_index']
                    train_abs_end = start_idx + train_segments[j]['end_index']
                    train_time_info = calculate_timestamps(start_date, train_abs_start, train_abs_end, FS)

                    # 按照要求的顺序排列JSON数据，添加label字段
                    json_data = {
                        "start_time": train_time_info["start_time"],
                        "end_time": train_time_info["end_time"],
                        "duration_seconds": train_time_info["duration_seconds"],
                        "start_index": train_time_info["start_index"],
                        "end_index": train_time_info["end_index"],
                        "event_id": i,
                        "train_index": j,
                        "type": "convergence",
                        "is_borrowed": train_info['is_borrowed'],
                        "borrowed_length": train_info['borrowed_length'],
                        "sampling_rate": FS,
                        "label": 1,  # 添加label字段，列车事件label=1
                        "original_start_index": int(start_idx),
                        "original_end_index": int(end_idx),
                        "sensor_id": sensor_id,
                        "sensor_name": sensor_name,
                        "data": t_data.tolist()
                    }
                    save_json(evt_json_dir, f"{fname_base}.json", json_data)

                    all_processed_trains.append({
                        'type': 'convergence',
                        'data_std': float(np.std(t_data)),
                        'event_id': i,
                        'train_id': j,
                        'time_info': train_time_info,
                        'label': 1
                    })

            else:
                # === 单列车事件处理 ===
                single_count += 1
                standardized = process_single_train_event(original_data, FS, TARGET_SAMPLES, SINGLE_TRAIN_MAX_SAMPLES)

                # 绘图 - 传入时间信息
                vis_path = os.path.join(evt_img_dir, f"single_event_{i:03d}.png")
                sep_train = [{'data': standardized, 'is_borrowed': False}]
                create_event_visualization(i, original_data, False, [], sep_train, FS, vis_path, original_time_info)

                fname_base = f"event_{i:03d}_single"

                # 按照要求的顺序排列JSON数据，添加label字段
                json_data = {
                    "start_time": original_time_info["start_time"],
                    "end_time": original_time_info["end_time"],
                    "duration_seconds": original_time_info["duration_seconds"],
                    "start_index": original_time_info["start_index"],
                    "end_index": original_time_info["end_index"],
                    "event_id": i,
                    "type": "single",
                    "is_borrowed": False,
                    "borrowed_length": 0,
                    "sampling_rate": FS,
                    "label": 1,  # 添加label字段，列车事件label=1
                    "sensor_id": sensor_id,
                    "sensor_name": sensor_name,
                    "data": standardized.tolist()
                }
                save_json(evt_json_dir, f"{fname_base}.json", json_data)

                all_processed_trains.append({
                    'type': 'single',
                    'data_std': float(np.std(standardized)),
                    'event_id': i,
                    'time_info': original_time_info,
                    'label': 1
                })

        except Exception as e:
            print(f"事件 {i} 处理错误: {e}")
            error_events.append(i)

    # 9. 生成列车事件汇总报告
    processing_time = (datetime.now() - start_time).total_seconds()
    stats_for_word = {
        'total_trains': len(all_processed_trains),
        'single_events': single_count,
        'convergence_events': conv_stats['count'],
        'trains_from_convergence': conv_stats['trains'],
        'borrowed_trains': conv_stats['borrowed'],
        'error_events': len(error_events)
    }
    create_word_report(
        event_run_dir, stats_for_word, conv_details, all_processed_trains, processing_time,
        evt_img_dir, evt_conv_img_dir, start_date, file_basename, threshold_info
    )

    print(f"\n📊 文件统计 {start_date}:")
    print(f"   单列车: {single_count} | 相汇: {conv_stats['count']}")
    print(f"   总列车数: {len(all_processed_trains)}")
    if norm_stats:
        print(f"   正常切片: {norm_stats['total']}")
    print(f"   自适应阈值: {enter_thr:.6f}")
    print(f"   耗时: {processing_time:.2f}s")

    return {
        "npy_file": file_basename,
        "date": start_date,
        "total_events": len(slices),
        "single_events": single_count,
        "convergence_events": conv_stats['count'],
        "total_trains": len(all_processed_trains),
        "processing_time": processing_time,
        "normal_slices": norm_stats["total"] if norm_stats else 0,
        "threshold": float(enter_thr),
        "event_run_dir": event_run_dir,
        "normal_run_dir": normal_run_dir
    }


def process_file_wrapper(json_path):
    try:
        return process_single_file(json_path)
    except Exception as e:
        print(f"Error: {e}")
        return None


# =========================
# 主流程 (批量处理)
# =========================
def main():
    freeze_support()  # Windows下如果是打包程序需要

    print(f"🚀 开始批量处理 (自适应阈值版本 - 稳定120左右)")
    print(f"   目标事件数量: {EXPECTED_MIN_EVENTS}-{EXPECTED_MAX_EVENTS} 个/天 (目标≈{EXPECTED_TARGET_EVENTS})")
    print(f"📁 输入目录: {INPUT_DIR}")
    print(f"📁 列车事件输出目录: {OUT_PARENT_DIR}")
    print(f"📁 正常振动输出目录: {NORMAL_VIB_PARENT_DIR}")

    # 查找所有1440拼接生成的JSON文件
    json_files = find_json_files(INPUT_DIR)

    if not json_files:
        print(f"❌ 未在 {INPUT_DIR} 找到JSON文件，请检查路径。")
        return

    file_count = len(json_files)
    print(f"📋 找到 {file_count} 个JSON文件准备处理。\n")

    # 固定使用10个CPU核心
    use_cores = 10
    print(f"🚀 启动并行处理池，固定使用核心数: {use_cores}")

    start_all = datetime.now()

    results = []

    # 使用 Pool 进行并行处理
    with Pool(processes=use_cores) as pool:
        # 使用map进行并行处理
        results = pool.map(process_file_wrapper, json_files)

    # 过滤掉None结果
    results = [r for r in results if r]

    success_count = len(results)
    fail_count = file_count - success_count

    end_all = datetime.now()
    total_time = (end_all - start_all).total_seconds()

    print(f"\n{'=' * 80}")
    print(f"🎉 批量处理完成!")
    print(f"{'=' * 80}")
    print(f"📊 总体统计:")
    print(f"   处理文件总数: {success_count}/{file_count}")
    print(f"   总处理时间: {total_time:.2f}秒")
    if results:
        avg_time = total_time / len(results)
        avg_events = np.mean([r['total_trains'] for r in results])
        std_events = np.std([r['total_trains'] for r in results])
        print(f"   平均每个文件处理时间: {avg_time:.2f}秒")
        print(f"   平均每天事件数: {avg_events:.1f} ± {std_events:.1f}")
        print(f"\n   各文件事件数:")
        for r in results:
            print(f"      {r['date']}: {r['total_trains']} 个 (阈值: {r['threshold']:.6f})")

    if results:
        total_train_events = sum(result['total_trains'] for result in results)
        total_normal_slices = sum(result['normal_slices'] for result in results)

        print(f"\n📋 标签统计:")
        print(f"   列车事件切片总数 (标签=1): {total_train_events}")
        print(f"   非列车事件切片总数 (标签=0): {total_normal_slices}")
        print(f"   总切片数: {total_train_events + total_normal_slices}")
        if total_train_events + total_normal_slices > 0:
            train_ratio = total_train_events / (total_train_events + total_normal_slices) * 100
            print(f"   列车事件占比: {train_ratio:.2f}%")
    print(f"{'=' * 80}")

    # 打印输出目录结构
    print(f"\n📁 输出目录结构:")
    print(f"{ROOT_OUTPUT_DIR}")
    print(f"└── {DEVICE_ID}列车事件切片/")
    print(f"    ├── Month_列车事件/")
    print(f"    │   ├── 2023-01-11/")
    print(f"    │   │   ├── images/")
    print(f"    │   │   ├── convergence_images/")
    print(f"    │   │   ├── slices_json/")
    print(f"    │   │   └── 处理报告.docx")
    print(f"    │   ├── 2023-01-12/")
    print(f"    │   └── ...")
    print(f"    └── Min_非列车事件/")
    print(f"        ├── 2023-01-11/")
    print(f"        │   ├── images/")
    print(f"        │   ├── json_slices/")
    print(f"        │   └── normal_vibration_report.docx")
    print(f"        ├── 2023-01-12/")
    print(f"        └── ...")

    print(f"\n✅ 全部完成")


if __name__ == "__main__":
    main()